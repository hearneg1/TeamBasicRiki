import base64

import markdown2
from docx import Document
from io import BytesIO
from pdfdocument.document import PDFDocument


def get_file_size(data):
    size_in_bytes = len(data)

    units = ['B', 'KB', 'MB', 'GB']
    factor = 1024
    size = size_in_bytes
    unit = units[0]

    if size < factor:
        return f"{size} {unit}"

    for next_unit in units[1:]:
        if size >= factor:
            size /= factor
            unit = next_unit
        else:
            break

    formatted_size = '{:.2f} {}'.format(size, unit)

    return formatted_size


class Converter(object):
    def __init__(self, page, filetype):
        self.page = page
        self.path = page.get_path()
        self.filetype = filetype

    def get_path(self):
        return self.path

    def convert_to_PDF(self):
        markdown_content = self.page.content
        pdf_buffer = BytesIO()
        pdf_content = PDFDocument(pdf_buffer)
        pdf_content.init_report()
        pdf_content.h1(self.page.title)
        pdf_content.p(markdown_content)
        pdf_content.generate()
        pdf_bytes = pdf_buffer.getvalue()
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        file_size = get_file_size(pdf_bytes)
        return pdf_base64, file_size

    def convert_to_TXT(self):
        txt_content = self.page.content.encode('utf-8')
        txt_base64 = base64.b64encode(txt_content).decode('utf-8')
        formatted_size = get_file_size(txt_content)
        return txt_base64, formatted_size

    def convert_to_HTML(self):
        html_content = markdown2.markdown(self.page.content)
        html_content_bytes = html_content.encode('utf-8')
        html_base64 = base64.b64encode(html_content_bytes).decode('utf-8')
        formatted_size = get_file_size(html_content_bytes)
        return html_base64, formatted_size

    def convert_to_DOCX(self):
        doc = Document()
        doc.add_paragraph(self.page.content)
        docx_content = BytesIO()
        doc.save(docx_content)
        docx_bytes = docx_content.getvalue()
        docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')
        file_size = get_file_size(docx_bytes)
        return docx_base64, file_size
